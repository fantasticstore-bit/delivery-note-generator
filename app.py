import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

st.title("📦 Delivery Note Generator - DOCX MODE")

excel_file = st.file_uploader("Upload Excel", type=["xlsx"])
generate = st.button("Generate")


# ==========================
# CREA DOCUMENTO
# ==========================
def create_doc(dn, data):

    doc = Document()

    row = data.iloc[0]

    # ==========================
    # HEADER
    # ==========================
    doc.add_paragraph(f"AFLEVERBON {dn}")
    doc.add_paragraph(f"Client: {row['CLIENT']}")

    # ==========================
    # CLIENT
    # ==========================
    doc.add_paragraph("")

    doc.add_paragraph(row["NAME"])
    doc.add_paragraph(row["STRASSE"])
    doc.add_paragraph(f"{row['PC']} {row['CITY']}")

    doc.add_paragraph("")

    # ==========================
    # TABLE
    # ==========================
    table = doc.add_table(rows=1, cols=5)

    headers = ["ARTIKEL", "AANTAL", "CONSUMENTEENHEID", "OMSCHRIJVING", "PRIJS"]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for _, r in data.iterrows():

        row_cells = table.add_row().cells

        row_cells[0].text = str(r["SKU"])
        row_cells[1].text = str(r["quantity_(bundles)"])
        row_cells[2].text = ""
        row_cells[3].text = ""  # descrizione opzionale
        row_cells[4].text = str(r["price per bundle"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ==========================
# MAIN
# ==========================
if generate and excel_file:

    df = pd.read_excel(excel_file)
    df.columns = df.columns.str.strip()

    for col in df.columns:
        if "bundles" in col.lower():
            df = df.rename(columns={col: "quantity_(bundles)"})

    grouped = df.groupby("DN")

    for dn, group in grouped:

        doc_file = create_doc(dn, group)

        st.download_button(
            f"Download {dn}",
            doc_file,
            file_name=f"{dn}.docx"
        )

    st.success("✅ DONE")
